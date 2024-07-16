from docx import Document

def replace_text_all(docxnya, jsonnya):
    def replace_text_in_header_footer(doc, search_text, replace_text):
        # Helper function to replace text in a paragraph while preserving formatting
        def replace_text_in_paragraph(paragraph, search_text, replace_text):
            for run in paragraph.runs:
                if search_text in run.text:
                    run.text = run.text.replace(search_text, replace_text)

        # Replace text in the header
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                replace_text_in_paragraph(paragraph, search_text, replace_text)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, search_text, replace_text)

            # Replace text in the footer
            footer = section.footer
            for paragraph in footer.paragraphs:
                replace_text_in_paragraph(paragraph, search_text, replace_text)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, search_text, replace_text)

        return doc


    def replace_and_print_table(docx_path, new_data, search_for):
        # Load the document
        doc = (docx_path)
        
        for table in doc.tables:
            # Update the data in the table
            for row in table.rows:
                label_cell = row.cells[0]
                # Check if the label matches like "Penyedia Jasa"
                if label_cell.text == search_for:
                    data_cell = row.cells[2]  # Assuming the data is in the third column
                    # Clear existing runs in the data cell
                    data_cell.text = ""
                    # Copy font properties from label_cell to data_cell
                    for paragraph in label_cell.paragraphs:
                        for run in paragraph.runs:
                            new_run = data_cell.paragraphs[0].add_run(new_data)
                            new_run.font.name = run.font.name
                            new_run.font.size = run.font.size
                            break  # We only need to copy the style once
                    break  # Exit after updating the first matching row

        # Save the document
        return doc


    ContractData = jsonnya

    # Call the function with the specific text to be replaced
    docxnya = replace_text_in_header_footer(docxnya, "CONTRACTOR_NAME", ContractData["CONTRACTOR_NAME"])

    # # Iterate over the key-value pairs and print them
    for key, value in ContractData.items():
        search_for = key
        replacement = value
        # if search_for == "REPORT_DATE_AND_TIME":
        #     print("Aktif!")
        docxnya = replace_and_print_table(docxnya, replacement, search_for)

    def extract_paragraphs_and_replace(docx_file, search_for, replacement):
        # Load the document
        document = (docx_file)
        
        # Iterate through each paragraph in the document
        for paragraph in document.paragraphs:
            # Check if the paragraph style is 'Normal'
            if paragraph.style.name == 'Normal':
                if search_for in paragraph.text:
                    # print("Found:", search_for)
                    # Replace the text
                    paragraph.text = paragraph.text.replace(search_for, replacement)
                    # print("Replaced with:", replacement)

        # Save the modified document
        return document

    def replace_exact_words_table(docx_path, new_data, search_for):
        # Load the document
        doc = (docx_path)
        
        for table in doc.tables:
            # Update the data in the table
            for row in table.rows:
                for cell in row.cells:
                    if search_for in cell.text:
                        # Replace the specific string with new_data
                        cell.text = cell.text.replace(search_for, new_data)
        
        # Save the document
        return doc


    # Example usage
    for key, value in ContractData.items():
        search_for = key
        replacement = value
        if search_for in ("TKDN Komitmen", "TKDN Self Assessment", "TKDN Hasil Verifikasi", "Penyedia Jasa"):
            continue
        docxnya = extract_paragraphs_and_replace(docxnya, search_for, replacement)

    for key, value in ContractData.items():
        search_for = key
        replacement = value
        if search_for not in ('REPORT_DATE_AND_TIME', 'COMMITMENT_TKDN_VALUE', 'VERIFICATOR_NAME'):
            continue
        docxnya = replace_exact_words_table(docxnya, replacement, search_for)

    return docxnya