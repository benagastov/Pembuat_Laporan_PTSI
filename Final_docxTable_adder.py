from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def adding_tables_all(docxnya, jsonnya, excelnya, dirname):
    def insert_table_under_paragraph(docx_file, last_paragraph, new_paragraph_text, excelpath, komitmen_TKDN, dirname):
        # Load the document
        document = docx_file
        
        # Variable to keep track of the current paragraph
        para_to_insert_after = None
        
        # Iterate through each paragraph in the document
        for paragraph in document.paragraphs:
            # print(paragraph.text)
            # Check if the paragraph text matches the last_paragraph
            if paragraph.text.startswith(last_paragraph):
                para_to_insert_after = paragraph
        
        # Check if we found the paragraph to insert after
        if para_to_insert_after is None:
            raise ValueError(f"The paragraph started with text '{last_paragraph}' was not found in the document.")
        
        # Add the new paragraph text
        new_para = document.add_paragraph(new_paragraph_text)
        new_para.style = document.styles['Normal']
        
        # Add a run to the new paragraph and set the text
        run = new_para.add_run()
        run.font.name = 'Segoe UI'
        run.font.size = Pt(10)
        
        # Insert the new paragraph after the specified paragraph
        para_to_insert_after._element.addnext(new_para._element)
        
        # Generate the table using the custom function
        if last_paragraph == 'Tabel 3':
            from final_SA_3_5 import process_documents
            table = process_documents(excelpath, docx_file, komitmen_TKDN, dirname)
        if last_paragraph == 'Tabel 4':
            from final_SA_3_5_strukturBiaya import process_documents
            table = process_documents(excelpath, docx_file, komitmen_TKDN, dirname)
        if last_paragraph == 'Tabel 5':
            from final_SA_3_3 import process_documents
            table = process_documents(excelpath, docx_file, dirname)
        if last_paragraph == 'Tabel 6':
            from final_SA_3_4 import process_documents
            table = process_documents(excelpath, docx_file, dirname)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        
        # Add the table after the new paragraph
        para_to_insert_after._element.addnext(table._element)
        
        # Save the document
        return document

    def append_mukadimah_table(docx_path, excelpath, docx_file, komitmen_TKDN):
        from docx.oxml.table import CT_Tbl
        from docx.shared import Cm
        from final_SA_3_5_mukadimah import process_documents
        # Load the document
        doc = docx_path
        new_table = process_documents(excelpath, docx_file, komitmen_TKDN)
        
        # Ensure new_table._element is an instance of CT_Tbl
        assert isinstance(new_table._element, CT_Tbl)
        
        # Get the first table's element
        first_table_element = doc.tables[0]._element
        
        # Get the parent of the first table's element
        parent_element = first_table_element.getparent()
        
        # Find the index of the first table's element within its parent element
        index_of_first_table = parent_element.index(first_table_element)
        
        # Insert a paragraph break after the first table
        paragraph = doc.add_paragraph()
        parent_element.insert(index_of_first_table + 1, paragraph._p)
        new_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in new_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # Insert the new table as a separate table after the paragraph break
        parent_element.insert(index_of_first_table + 2, new_table._element)

        # Save the document
        return doc


    docx_file = docxnya
    last_paragraphs = ['Tabel 3', 'Tabel 4', 'Tabel 5','Tabel 6']
    excelpath = excelnya

    ContractData = jsonnya

    komitmen_TKDN = ContractData["COMMITMENT_TKDN_VALUE"]

    append_mukadimah_table(docx_file, excelpath, docx_file, komitmen_TKDN)
    for last_paragraph in last_paragraphs:
        insert_table_under_paragraph(docx_file, last_paragraph, None, excelpath, komitmen_TKDN, dirname)

    return docxnya