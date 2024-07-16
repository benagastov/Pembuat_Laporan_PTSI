from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE

def process_documents(excelpath, wordpath, dirname):
    # Function to set cell borders
    def set_cell_border(cell, **kwargs):
        """
        Set cell's border
        Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        for edge in ('start', 'top', 'end', 'bottom'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = tcPr.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcPr.append(element)
                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    border_style = {
            "top": {"sz": 5, "val": "single", "color": "B4C6E7"},
            "bottom": {"sz": 5, "val": "single", "color": "B4C6E7"},
            "start": {"sz": 5, "val": "single", "color": "B4C6E7"},
            "end": {"sz": 5, "val": "single", "color": "B4C6E7"}
        }

    # Function to set cell shading (background color)
    def set_cell_shading(cell, fill, color=None):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill)
        if color:
            shading_elm.set(qn('w:color'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Load the workbook and select the specified sheet
    workbook = load_workbook(excelpath)
    sheet = workbook['Form 3.3']

    # Find the row and column index for the cell with the value "Uraian"
    uraian_row_index = uraian_column_index = None
    for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
        for cell_index, cell in enumerate(row, start=1):
            # print(f"Checking cell: {cell}")  # Debug print statement
            if cell == "Uraian":
                uraian_row_index = row_index
                uraian_column_index = cell_index+1
                break
        if uraian_row_index and uraian_column_index:
            break

    # Check if the "Uraian" cell was found
    all_rows_in_table = []
    stop_adding_rows = False  # Flag to indicate when to stop adding rows

    if uraian_row_index is None or uraian_column_index is None:
        print("The cell with value 'Uraian' was not found in the sheet.")
    else:
        # Extract all rows below the "Uraian" cell and in the same column, excluding empty cells
        for row in sheet.iter_rows(min_row=uraian_row_index+1, min_col=uraian_column_index, max_col=uraian_column_index):
            if stop_adding_rows:  # Check if we need to stop adding rows
                break
            for cell in row:  # Iterate over cells in the row
                if cell.border.top.border_style == "medium":
                    stop_adding_rows = True  # Set the flag to stop adding rows
                    break
                elif cell.value is not None and cell.value != "None":  # Check if the cell is not empty
                    all_rows_in_table.append(cell.row)

        # print("Uraian:", all_rows_in_table)

        

        # Load a Document
        doc = Document()

        # Add a table to the document
        table = doc.add_table(rows=1, cols=6)

        # Define a new paragraph style with no space after if not already defined
        style = doc.styles['Normal']
        style = doc.styles.add_style('NoSpaceAfter', WD_STYLE_TYPE.PARAGRAPH)
        style.paragraph_format.space_after = Pt(0)

        # Set the font size for the 'NoSpaceAfter' style
        font = style.font
        font.name = 'Segoe UI'
        font.size = Pt(9)

        # Apply the 'NoSpaceAfter' style to all paragraphs in the table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles['NoSpaceAfter']

        # Set the headers and add the rows to the table
        import locale

        def format_percentage(value):
            # Format the value as a decimal with two digits after the comma
            return '0,00' if value in (None, 0) else f"{value:.2f}".replace(".", ",")

        # Calculate the value for "Header 5" based on the Excel formula
        form_3_5_sheet = workbook['Form 3.5']
        g22_value = form_3_5_sheet['G22'].value

        # Initialize the list to store formatted values for "Header 5"
        baris_struktur_biaya_formatted = []
        import formulas
        import os.path as path

        # Reload the workbook with openpyxl to access the evaluated values
        workbook_calc = load_workbook(excelpath, data_only=True)
        sheet_calc = workbook_calc['Form 3.3']

        # Initialize total_calculated_value to 0
        total_calculated_value = 0

        # Define g22_value (you'll need to replace this with the actual value from cell G22)
        form_3_5_sheet = workbook_calc['Form 3.5']
        g22_value = form_3_5_sheet['G22'].value

        # Create an empty list to store formatted values
        baris_struktur_biaya_formatted = []
        # print("nilai G", g22_value)

        # Iterate over the specified rows and columns
        for row in sheet_calc.iter_rows(min_row=min(all_rows_in_table), max_row=max(all_rows_in_table), min_col=12, max_col=12):
            for cell in row:
                l_value = cell.value if cell.value is not None else 0
                l_value *= 100
                # print("nilai L", l_value)
                
                # Calculate the value for "Header 5"
                calculated_value = l_value / g22_value if g22_value else 0
                total_calculated_value += calculated_value
                
                # Format and update the "Header 5" cell
                formatted_value = format_percentage(calculated_value)
                cell.value = formatted_value
                baris_struktur_biaya_formatted.append(formatted_value)


        # Calculate and format the total for "Header 5"
        total_struktur_biaya_formatted = format_percentage(total_calculated_value)


        # Assuming 'all_rows_in_table' is a list of row indices and 'sheet' is a worksheet object
        # The header row is set separately and not included in the loop
        table_header = ["No.","Uraian Jabatan", "Nama Perusahaan", "Kewarganegaraan", "TKDN (%)", 'Struktur Biaya (%)']
        header_cells = table.rows[0].cells
        for index, header in enumerate(table_header):
            header_cells[index].text = header
            header_cells[index].paragraphs[0].style = style
            header_cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(header_cells[index], **border_style)
            set_cell_shading(header_cells[index], fill='31849B')
            header_cells[index].paragraphs[0].runs[0].font.bold = True
            header_cells[index].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # Start adding data from the second row
        def format_percentage(value):
            # Format the value as a decimal with two digits after the comma
            return '0,00' if value in (None, 0) else f"{value:.2f}".replace(".", ",")
        
        for row_index, row_data in enumerate([['',
                str(sheet.cell(row=row_index, column=uraian_column_index).value),
                "Uniqlo Indonesia",
                str(sheet.cell(row=row_index, column=5).value),
                format_percentage(sheet.cell(row=row_index, column=6).value*100)] for row_index in all_rows_in_table], start=2):
            row_cells = table.add_row().cells
            # print(row_data)
            for cell_index, cell in enumerate(row_cells):
                if cell_index < len(row_data):
                    cell.text = row_data[cell_index]
                elif cell_index == 5:  # Insert values from baris_struktur_biaya into the "Struktur Biaya" column
                    if row_index - 2 < len(baris_struktur_biaya_formatted):
                        cell.text = str(baris_struktur_biaya_formatted[row_index - 2])
                    else:
                        cell.text = "0,00"   # Placeholder text or handle as needed
                cell.paragraphs[0].style = style
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_border(cell, **border_style)

                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_border(cell, **border_style)

        
        # Initialize the list to keep track of rows to remove
        rows_to_remove = []

        # Iterate over the cells in the "Struktur Biaya" column (assuming it's the 6th column with index 5)
        for idx, cell in enumerate(table.columns[5].cells):
            cell.width = Cm(0.99)  # Set the width of the cell
            # Check if the cell text represents a value not greater than "0,00"
            if cell.text.strip() in ('-', 'None', '0,00'):
                rows_to_remove.append(idx)  # Add the row index to the list of rows to remove

        # Sort the list of rows to remove in reverse order to avoid shifting indices
        rows_to_remove.sort(reverse=True)

        # Remove the rows from the table
        for idx in rows_to_remove:
            tr = table.rows[idx]._tr  # Access the underlying XML element of the row
            tr.getparent().remove(tr)  # Remove the row from its parent element

        # Define the background color
        bg_color = parse_xml(r'<w:shd {} w:fill="31849B"/>'.format(nsdecls('w')))

        # Add numbering to the first cell of each row (excluding the header)
        for idx, row in enumerate(table.rows[1:], start=1):  # Skip the header row
            row.cells[0].text = str(idx)
            row.cells[0].paragraphs[0].style = style

        # Add a new row at the end of the table
        total_row = table.add_row()

        # Set the background color of the first cell
        total_row.cells[0]._tc.get_or_add_tcPr().append(bg_color)

        # Merge cells from the second to the fourth column
        merged_cells = total_row.cells[1].merge(total_row.cells[4])

        # Set the merged cells value to "Total"
        merged_cells.text = "Total"
        merged_cells.paragraphs[0].style = style
        merged_cells.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        merged_cells.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        merged_cells.paragraphs[0].runs[0].font.bold = True

        # Calculate the sum of all numbers in the fifth column and set it as the value of the fifth cell
        total_row.cells[5].text = str(total_struktur_biaya_formatted)
        for cell in total_row.cells:
            set_cell_border(cell, **border_style)
        total_row.cells[5].paragraphs[0].style = style
        total_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_row.cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        total_row.cells[5].paragraphs[0].runs[0].font.bold = True


        # Set the width of the "TKDN" column
        for row in table.rows:
            row.cells[0].width = Cm(0.5)
        for row in table.rows:
            row.cells[1].width = Cm(5.48)
        for row in table.rows:
            row.cells[2].width = Cm(4)
        for row in table.rows:
            row.cells[4].width = Cm(1.5)
        for row in table.rows:
            row.cells[5].width = Cm(1.5)

        # Return table
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            # print(row_text)
        
        return table


# process_documents(excelpath, wordpath)