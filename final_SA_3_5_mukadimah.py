from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
import formulas
import os.path as path

def process_documents(excelpath, wordpath, komitmen_TKDN):
    # Function to set cell borders
    def set_cell_border(cell, **kwargs):
        """
        Set cell's border
        Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        for edge in ('start', 'top', 'end', 'bottom'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = tcPr.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcPr.append(element)
                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    border_style = {
            "top": {"sz": 5, "val": "single", "color": "B4C6E7"},
            "bottom": {"sz": 5, "val": "single", "color": "B4C6E7"},
            "start": {"sz": 5, "val": "single", "color": "B4C6E7"},
            "end": {"sz": 5, "val": "single", "color": "B4C6E7"}
        }

    # Function to set cell shading (background color)
    def set_cell_shading(cell, fill, color=None):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill)
        if color:
            shading_elm.set(qn('w:color'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Load the workbook and select the specified sheet
    workbook = load_workbook(excelpath)
    sheet = workbook['Form 3.5']

    # Find the row and column index for the cell with the value "Uraian"
    uraian_row_index = uraian_column_index = None
    for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
        for cell_index, cell in enumerate(row, start=1):
            # print(f"Checking cell: {cell}")  # Debug print statement
            if cell == "Jasa":
                uraian_row_index = row_index
                uraian_column_index = cell_index
                break
        if uraian_row_index and uraian_column_index:
            break

    # Check if the "Uraian" cell was found
    all_rows_in_table = []
    stop_adding_rows = False  # Flag to indicate when to stop adding rows

    if uraian_row_index is None or uraian_column_index is None:
        print("The cell with value 'Uraian' was not found in the sheet.")
    else:
        # Extract all rows below the "Uraian" cell and in the same column, excluding empty cells
        for row in sheet.iter_rows(min_row=uraian_row_index+1, min_col=uraian_column_index, max_col=uraian_column_index):
            if stop_adding_rows:  # Check if we need to stop adding rows
                break
            for cell in row:  # Iterate over cells in the row
                if cell.border.top.border_style == "medium":
                    stop_adding_rows = True  # Set the flag to stop adding rows
                    break
                elif cell.value is not None and cell.value != "None":  # Check if the cell is not empty
                    all_rows_in_table.append(cell.row)

        # print("Uraian:", all_rows_in_table)
        

        # Create a new Document
        doc = Document()

        # Add a table to the document with two rows for headers and enough columns
        table = doc.add_table(rows=2, cols=8)

        # Access the first row which will be used for the merged header "TKDN (%)"
        first_row_cells = table.rows[0].cells

        # Merge the cells for "TKDN (%)" across the columns for sub-headers
        tkdn_cell = (first_row_cells[2]).merge(first_row_cells[3]).merge(first_row_cells[4])
        tkdn_cell.text = 'TKDN (%)'

        # Merge the cells for "TKDN (%)" across the columns for sub-headers
        tkdn_cell = (first_row_cells[5]).merge(first_row_cells[6]).merge(first_row_cells[7])
        tkdn_cell.text = 'Struktur Biaya (%)'

        # Access the second row to add sub-headers
        second_row_cells = table.rows[1].cells

        # Define a new paragraph style with no space after if not already defined
        style = doc.styles['Normal']
        style = doc.styles.add_style('NoSpaceAfter', WD_STYLE_TYPE.PARAGRAPH)
        style.paragraph_format.space_after = Pt(0)

        # Set the font size for the 'NoSpaceAfter' style
        font = style.font
        font.name = 'Segoe UI'
        font.size = Pt(9)

        # Apply the 'NoSpaceAfter' style to all paragraphs in the table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles['NoSpaceAfter']


        # Calculate the value for "Header 5" based on the Excel formula
        form_3_5_sheet = workbook['Form 3.5']
        g22_value = form_3_5_sheet['G22'].value

        # Initialize the list to store formatted values for "Header 5"
        baris_struktur_biaya_formatted = []

        # Reload the workbook with openpyxl to access the evaluated values
        workbook_calc = load_workbook(excelpath, data_only=True)
        sheet_calc = workbook_calc['Form 3.5']
        sheet_calcSA = workbook_calc['SA']

        def struktur_biaya(value, rownya, columnnya):
            try:
                totaltotalKDN = (sheet_calc.cell(row=22, column=columnnya-1).value)
                totalKDN = (sheet_calc.cell(row=rownya, column=columnnya-1).value*100)/totaltotalKDN
            except:
                totalKDN = 0
            return format_percentage(totalKDN)


        def format_cell(cell, text, style, border_style, boldness):
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.style = style
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.runs[0].font.bold = boldness
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, **border_style)


        # Initialize total_calculated_value to 0
        total_calculated_value = 0

        # Assuming 'all_rows_in_table' is a list of row indices and 'sheet' is a worksheet object
        # The header row is set separately and not included in the loop
        table_header = ["No.","Komponen Biaya", "Komitmen", "Self Assessment", "Verifikasi", "Komitmen", "Self Assessment", "Verifikasi"]
        for index, header in enumerate(table_header):
            second_row_cells[index].text = header
            if header not in ("Komitmen", "Self Assessment", "Verifikasi"):
                # Convert the tuple to a list, merge the cells, then convert back to a tuple
                cells_list = list(second_row_cells)
                cells_list[index] = (first_row_cells[index]).merge(cells_list[index])
                second_row_cells = tuple(cells_list)
                second_row_cells[index].text = header
            second_row_cells[index].paragraphs[0].style = style
            second_row_cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            second_row_cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(second_row_cells[index], fill='31849B')
            second_row_cells[index].paragraphs[0].runs[0].font.bold = True
            second_row_cells[index].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Assuming you have the second_row_cells styles set up as desired
        # Loop through each cell in the first row and apply the styles from the second row
        for first_cell, second_cell in zip(first_row_cells, second_row_cells):
            for paragraph in first_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = second_cell.paragraphs[0].runs[0].font.name
                    run.font.size = second_cell.paragraphs[0].runs[0].font.size
                    run.font.bold = second_cell.paragraphs[0].runs[0].font.bold
                    run.font.color.rgb = second_cell.paragraphs[0].runs[0].font.color.rgb
            
            # Set cell alignment
            first_cell.paragraphs[0].alignment = second_cell.paragraphs[0].alignment
            first_cell.vertical_alignment = second_cell.vertical_alignment
            
            # Apply border style
            set_cell_border(first_cell, **border_style)
            set_cell_border(second_cell, **border_style)
            
            # Apply cell shading (background color)
            # Assuming the second cell has a shading element, we can use it to set the background for the first cell
            second_cell_shading = second_cell._tc.get_or_add_tcPr().find(qn('w:shd'))
            fill = second_cell_shading.get(qn('w:fill'))
            set_cell_shading(first_cell, fill=fill)

        # Start adding data from the second row
        def format_percentage(value):
            # Format the value as a decimal with two digits after the comma
            return '0,00' if value in (None, 0) else f"{value:.2f}".replace(".", ",")
        def format_percentage_in_float(value):
            # Remove the percentage sign and convert to float
            value = float(value.strip('%'))
            # Format the value as a decimal with two digits after the comma
            return float(0) if value == 0 else float(value)
        
        import re

        def remove_roman_numerals(input_string):
            # Regex pattern to match Roman numerals followed by a period at the start of a string
            pattern = r'(?i)^(M{0,4}(CM|CD|D?C{0,3})(XC|XL|L?X{0,3})(IX|IV|V?I{0,3})\.)\s+'
            # Replace the matched pattern with an empty string
            return re.sub(pattern, '', input_string)


        sheet2 = workbook_calc['SA']
        for row_index, row_data in enumerate([["",
                remove_roman_numerals(str(sheet.cell(row=row_index, column=uraian_column_index).value)),
                "",
                format_percentage(sheet2.cell(row=row_index, column=8).value*100),
                format_percentage(sheet_calc.cell(row=row_index, column=8).value*100),
                "",
                struktur_biaya(sheet2.cell(row=row_index, column=8).value, row_index,sheet_calc.cell(row=row_index, column=8).column),
                struktur_biaya(sheet_calc.cell(row=row_index, column=8).value, row_index,sheet_calc.cell(row=row_index, column=8).column), ''] for row_index in all_rows_in_table], start=2):
            row_cells = table.add_row().cells
            # print(row_data)
            for cell_index, cell in enumerate(row_cells):
                if cell_index < len(row_data):
                    format_cell(cell, row_data[cell_index], style, border_style, False)
                elif cell_index == 5:  # Insert values from baris_struktur_biaya into the "Struktur Biaya" column
                    if row_index - 2 < len(baris_struktur_biaya_formatted):
                        format_cell(cell, str(baris_struktur_biaya_formatted[row_index - 2]), style, border_style, False)
                    else:
                        format_cell(cell, "0,00", style, border_style, False)   # Placeholder text or handle as needed
        

        # Define the background color
        bg_color = parse_xml(r'<w:shd {} w:fill="31849B"/>'.format(nsdecls('w')))

        # Add numbering to the first cell of each row (excluding the header)
        for idx, row in enumerate(table.rows[2:], start=1):  # Skip the header row
            row.cells[0].text = str(idx)
            row.cells[0].paragraphs[0].style = style

        # Add a new row at the end of the table
        total_row = table.add_row()

        # Set the background color of the first cell
        total_row.cells[0]._tc.get_or_add_tcPr().append(bg_color)

        # Set the merged cells value to "Total"
        # total_row.cells[1].text = "Total"
        # total_row.cells[1].paragraphs[0].style = style
        # total_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # total_row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # total_row.cells[1].paragraphs[0].runs[0].font.bold = True
        format_cell(total_row.cells[1], "Total", style, border_style, True)


        # Define g22_value (you'll need to replace this with the actual value from cell G22)
        form_3_5_sheet = workbook_calc['Form 3.5']
        g22_value = form_3_5_sheet['G22'].value

        # Create an empty list to store formatted values
        totalTKDN_Verif = []
        total_calculated_Verifvalue = 0
        
        # Create an empty list to store formatted values
        totalStrukturBiayaTKDN_Verif = []
        total_calculated_StrukturVerifvalue = 0
        # Iterate over the specified rows and columns
        for row in sheet_calc.iter_rows(min_row=min(all_rows_in_table), max_row=max(all_rows_in_table), min_col=8 ,max_col=8):
            for cell in row:
                s_value = (cell.value)
                l_value = struktur_biaya(cell.value, cell.row, cell.column)
                
                # Calculate the value for "Header 5"
                calculated_value = float(l_value.replace(',','.'))
                calculateds_value = float(s_value)

                calculateds_value *= 100
                
                total_calculated_StrukturVerifvalue += calculated_value
                total_calculated_Verifvalue += calculateds_value
                
                # Format and update the "Header 5" cell
                formatted_value = format_percentage(calculated_value)
                cell.value = formatted_value
                totalStrukturBiayaTKDN_Verif.append(formatted_value)

                # Format and update the "Header 5" cell
                formatteds_value = format_percentage(calculateds_value)
                cell.value = formatteds_value
                totalTKDN_Verif.append(formatteds_value)
        
        # Create an empty list to store formatted values
        totalStrukturTKDN_SA = []
        total_calculated_StrukturSAvalue = 0
        # Create an empty list to store formatted values
        totalTKDN_SA = []
        total_calculated_SAvalue = 0

        # Iterate over the specified rows and columns
        for row in sheet2.iter_rows(min_row=min(all_rows_in_table), max_row=max(all_rows_in_table), min_col=8 ,max_col=8):
            for cell in row:
                l_value = struktur_biaya(cell.value, cell.row, cell.column)
                s_value = (cell.value) 

                # Calculate the value for "Header 5"
                calculated_value = float(l_value.replace(',','.'))
                total_calculated_StrukturSAvalue += calculated_value
                calculateds_value = float(s_value)
                
                calculateds_value *= 100
                total_calculated_SAvalue += calculateds_value  
                # Format and update the "Header 5" cell
                formatted_value = format_percentage(calculated_value)
                cell.value = formatted_value
                totalStrukturTKDN_SA.append(formatted_value)

                # Format and update the "Header 5" cell
                formatteds_value = format_percentage(calculateds_value)
                cell.value = formatteds_value
                totalTKDN_SA.append(formatteds_value)


        # Calculate and format the total for "Header 5"
        totalTKDN_SA_formatted = format_percentage(total_calculated_SAvalue)

        # Calculate and format the total for "Header 5"
        totalTKDN_Verif_formatted = format_percentage(total_calculated_Verifvalue)

        # Calculate and format the total for "Header 5"
        totalStrukturTKDN_SA_formatted = format_percentage(total_calculated_StrukturSAvalue)
        # Calculate and format the total for "Header 5"
        totalStrukturBiayaTKDN_Verif_formatted = format_percentage(total_calculated_StrukturVerifvalue)

        # Now you can call this function for each cell that needs formatting
        format_cell(total_row.cells[2], str(format_percentage(format_percentage_in_float(komitmen_TKDN))), style, border_style, True)
        format_cell(total_row.cells[3], str(totalTKDN_SA_formatted), style, border_style, True)
        format_cell(total_row.cells[4], str(totalTKDN_Verif_formatted), style, border_style, True)
        format_cell(total_row.cells[5], str(100), style, border_style, True)
        format_cell(total_row.cells[6], str(totalStrukturTKDN_SA_formatted), style, border_style, True)
        format_cell(total_row.cells[7], str(totalStrukturBiayaTKDN_Verif_formatted), style, border_style, True)



        # Set the width of each column
        ukuran = [0.5, 5.48, 4, 1.5, 2.07]
        for semuaKolom, centimeter in zip(range(0,6), ukuran):
            for row in table.rows:
                row.cells[semuaKolom].width = Cm(centimeter)

        # Check if the cell is empty or contains 'None' and set the background color to black
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if (not cell.text.strip() or cell.text.strip().lower() == 'none'):
                        set_cell_shading(cell,fill='31849B')


        # Return table
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            # print(row_text)
        
        return table

excelpath = 'Excels.xlsx'
wordpath = 'testdocxpy.docx'
# process_documents(excelpath, wordpath)