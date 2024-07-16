from docx import Document
from docx.shared import Inches

def add_logo_toHeaders(doc, img):

    # Flag to check if the image has been added
    image_added = False

    # Iterate through each section in the document
    for section in doc.sections:
        # Access the header of the current section
        header = section.header

        # Check if the header is not None and the image hasn't been added yet
        if header and not image_added:
            # Assuming the table is the first table in the header
            # and you want to add the image to the second column of the first row
            table = header.tables[0]
            cell = table.rows[0].cells[1]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture('./logoKontraktor.png', width=Inches(1))  # Adjust the width as needed

            # Set the flag to True after adding the image
            image_added = True

    # Save the document
    return doc
